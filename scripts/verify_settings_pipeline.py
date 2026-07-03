#!/usr/bin/env python3
"""
E2E Verification Test Suite for PDF Utility Platform Settings Pipeline.

Tests the following critical paths:
1. image_convert.py  - pdf2jpg with PNG format (was broken, now fixed)
2. image_convert.py  - pdf2jpg with JPG format (regression check)
3. doc_convert.py    - word2pdf with landscape orientation
4. doc_convert.py    - pdf2word with mode=flowing
5. doc_convert.py    - pdf2excel with singleSheet=True
6. doc_convert.py    - pdf2ppt with slideSize=4:3
7. doc_convert.py    - html2pdf with landscape orientation
8. basic_manipulation.py - merge + split + rotate (regression check)

Run from project root:
  python scripts/verify_settings_pipeline.py
"""

import os
import sys
import json
import subprocess
import zipfile
import tempfile
import shutil

# Detect python-engine venv
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
MODULES_DIR = os.path.join(PROJECT_ROOT, "python-engine", "modules")

WIN_VENV = os.path.join(PROJECT_ROOT, "python-engine", "venv", "Scripts", "python.exe")
NIX_VENV = os.path.join(PROJECT_ROOT, "python-engine", "venv", "bin", "python")

if os.path.exists(WIN_VENV):
    PYTHON = WIN_VENV
elif os.path.exists(NIX_VENV):
    PYTHON = NIX_VENV
else:
    PYTHON = sys.executable  # fallback

print(f"[INFO] Using Python: {PYTHON}")
print(f"[INFO] Modules Dir:  {MODULES_DIR}")
print()

PASS = 0
FAIL = 0

def run(script, payload):
    """Send payload as JSON to the Python module via stdin."""
    script_path = os.path.join(MODULES_DIR, script)
    result = subprocess.run(
        [PYTHON, script_path],
        input=json.dumps(payload),
        capture_output=True,
        text=True,
        timeout=60
    )
    try:
        data = json.loads(result.stdout.strip())
    except Exception:
        data = {"success": False, "error": f"Bad stdout: {result.stdout[:300]} | stderr: {result.stderr[:300]}"}
    return data

def check(test_name, condition, detail=""):
    global PASS, FAIL
    if condition:
        print(f"  PASS: {test_name}")
        PASS += 1
    else:
        print(f"  FAIL: {test_name}")
        if detail:
            print(f"     Detail: {detail}")
        FAIL += 1

def section(title):
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")

# -----------------------------------------------------------------
# Fixtures: create minimal test assets
# -----------------------------------------------------------------
tmp_dir = tempfile.mkdtemp(prefix="pdf_verify_")

# Minimal valid 1-page PDF (PDF spec hand-crafted)
MINI_PDF = os.path.join(tmp_dir, "test_input.pdf")
mini_pdf_bytes = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>
stream
BT /F1 12 Tf 100 700 Td (Hello PDF World) Tj ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000274 00000 n 
0000000352 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
449
%%EOF"""
with open(MINI_PDF, "wb") as f:
    f.write(mini_pdf_bytes)

# Minimal HTML file
HTML_FILE = os.path.join(tmp_dir, "test.html")
with open(HTML_FILE, "w", encoding="utf-8") as f:
    f.write("<html><body><h1>Hello</h1><p>Test paragraph for settings pipeline verification.</p></body></html>")

# Minimal DOCX (we use the python-docx library to create one)
DOCX_FILE = os.path.join(tmp_dir, "test.docx")
try:
    import docx
    doc = docx.Document()
    doc.add_heading("Settings Pipeline Verification", 0)
    doc.add_paragraph("This document was auto-generated for E2E testing.")
    doc.save(DOCX_FILE)
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx not available in current env -- skipping word2pdf test.")

# Minimal XLSX
XLSX_FILE = os.path.join(tmp_dir, "test.xlsx")
try:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 42])
    wb.save(XLSX_FILE)
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False
    print("[WARN] openpyxl not available in current env -- skipping excel2pdf test.")

# -----------------------------------------------------------------
# TEST 1: image_convert.py  pdf2jpg -> JPG format
# -----------------------------------------------------------------
section("TEST 1: pdf2jpg with JPG format (regression check)")
out_dir_1 = os.path.join(tmp_dir, "pdf2jpg_jpg")
os.makedirs(out_dir_1, exist_ok=True)
try:
    r = run("image_convert.py", {
        "action": "pdf2jpg",
        "file": MINI_PDF,
        "output_dir": out_dir_1,
        "format": "jpg",
        "dpi": 72,
        "quality": 70
    })
    check("pdf2jpg call succeeded", r.get("success"), r.get("error"))
    if r.get("success"):
        zip_path = r.get("output", "")
        check("Output zip exists", os.path.exists(zip_path), zip_path)
        if os.path.exists(zip_path):
            with zipfile.ZipFile(zip_path) as zf:
                names = zf.namelist()
            check("Zip contains .jpg file", any(n.endswith(".jpg") for n in names), str(names))
            check("Zip contains NO .png file", not any(n.endswith(".png") for n in names), str(names))
except subprocess.TimeoutExpired:
    check("pdf2jpg JPG - no timeout", False, "Timed out (likely poppler missing on this machine)")

# -----------------------------------------------------------------
# TEST 2: image_convert.py  pdf2jpg -> PNG format  (THE KEY FIX)
# -----------------------------------------------------------------
section("TEST 2: pdf2jpg with PNG format (the critical bug fix)")
out_dir_2 = os.path.join(tmp_dir, "pdf2jpg_png")
os.makedirs(out_dir_2, exist_ok=True)
try:
    r = run("image_convert.py", {
        "action": "pdf2jpg",
        "file": MINI_PDF,
        "output_dir": out_dir_2,
        "format": "png",
        "dpi": 72,
        "quality": 90
    })
    check("pdf2jpg (PNG) call succeeded", r.get("success"), r.get("error"))
    if r.get("success"):
        zip_path = r.get("output", "")
        check("Output zip exists", os.path.exists(zip_path), zip_path)
        if os.path.exists(zip_path):
            with zipfile.ZipFile(zip_path) as zf:
                names = zf.namelist()
            check("Zip contains .png file - FORMAT FIX VERIFIED", any(n.endswith(".png") for n in names), str(names))
            check("Zip contains NO .jpg file - FORMAT FIX VERIFIED", not any(n.endswith(".jpg") for n in names), str(names))
except subprocess.TimeoutExpired:
    check("pdf2jpg PNG - no timeout", False, "Timed out (likely poppler missing on this machine)")

# -----------------------------------------------------------------
# TEST 3: doc_convert.py  html2pdf -> default (portrait)
# -----------------------------------------------------------------
section("TEST 3: html2pdf with default orientation (portrait)")
out_html_portrait = os.path.join(tmp_dir, "html_portrait.pdf")
r = run("doc_convert.py", {
    "action": "html2pdf",
    "file": HTML_FILE,
    "output": out_html_portrait,
    "settings": {"orientation": "auto", "layoutMode": "fit"}
})
check("html2pdf (portrait) call succeeded", r.get("success"), r.get("error"))
check("html2pdf (portrait) output file exists", os.path.exists(out_html_portrait))
if os.path.exists(out_html_portrait):
    check("html2pdf (portrait) file is non-empty", os.path.getsize(out_html_portrait) > 0)

# -----------------------------------------------------------------
# TEST 4: doc_convert.py  html2pdf -> landscape orientation
# -----------------------------------------------------------------
section("TEST 4: html2pdf with LANDSCAPE orientation (settings wire test)")
out_html_landscape = os.path.join(tmp_dir, "html_landscape.pdf")
r = run("doc_convert.py", {
    "action": "html2pdf",
    "file": HTML_FILE,
    "output": out_html_landscape,
    "settings": {"orientation": "landscape", "layoutMode": "fit"}
})
check("html2pdf (landscape) call succeeded", r.get("success"), r.get("error"))
check("html2pdf (landscape) output file exists", os.path.exists(out_html_landscape))

if os.path.exists(out_html_portrait) and os.path.exists(out_html_landscape):
    landscape_size = os.path.getsize(out_html_landscape)
    portrait_size = os.path.getsize(out_html_portrait)
    check("Landscape PDF is valid (non-zero size)", landscape_size > 0)
    print(f"     Portrait size: {portrait_size}B  |  Landscape size: {landscape_size}B")

# -----------------------------------------------------------------
# TEST 5: doc_convert.py  word2pdf (if docx available)
# -----------------------------------------------------------------
section("TEST 5: word2pdf with landscape orientation")
if HAS_DOCX:
    out_word_pdf = os.path.join(tmp_dir, "word_landscape.pdf")
    r = run("doc_convert.py", {
        "action": "word2pdf",
        "file": DOCX_FILE,
        "output": out_word_pdf,
        "settings": {"orientation": "landscape", "layoutMode": "fit"}
    })
    check("word2pdf (landscape) call succeeded", r.get("success"), r.get("error"))
    check("word2pdf (landscape) output file exists", os.path.exists(out_word_pdf))
    if os.path.exists(out_word_pdf):
        check("word2pdf (landscape) file is non-empty", os.path.getsize(out_word_pdf) > 0)
else:
    print("  [SKIP] python-docx not in env")

# -----------------------------------------------------------------
# TEST 6: doc_convert.py  excel2pdf (if openpyxl available)
# -----------------------------------------------------------------
section("TEST 6: excel2pdf with portrait orientation")
if HAS_XLSX:
    out_excel_pdf = os.path.join(tmp_dir, "excel_portrait.pdf")
    r = run("doc_convert.py", {
        "action": "excel2pdf",
        "file": XLSX_FILE,
        "output": out_excel_pdf,
        "settings": {"orientation": "portrait", "layoutMode": "original"}
    })
    check("excel2pdf (portrait) call succeeded", r.get("success"), r.get("error"))
    check("excel2pdf (portrait) output file exists", os.path.exists(out_excel_pdf))
    if os.path.exists(out_excel_pdf):
        check("excel2pdf (portrait) file is non-empty", os.path.getsize(out_excel_pdf) > 0)
else:
    print("  [SKIP] openpyxl not in env")

# -----------------------------------------------------------------
# TEST 7: doc_convert.py  pdf2word
# -----------------------------------------------------------------
section("TEST 7: pdf2word with mode=flowing")
out_word = os.path.join(tmp_dir, "out.docx")
r = run("doc_convert.py", {
    "action": "pdf2word",
    "file": MINI_PDF,
    "output": out_word,
    "settings": {"mode": "flowing", "ocr": False}
})
check("pdf2word call succeeded", r.get("success"), r.get("error"))
check("pdf2word output .docx exists", os.path.exists(out_word))
if os.path.exists(out_word):
    check("pdf2word output is non-empty", os.path.getsize(out_word) > 0)

# -----------------------------------------------------------------
# TEST 8: doc_convert.py  pdf2excel
# -----------------------------------------------------------------
section("TEST 8: pdf2excel with singleSheet=false")
out_excel = os.path.join(tmp_dir, "out.xlsx")
r = run("doc_convert.py", {
    "action": "pdf2excel",
    "file": MINI_PDF,
    "output": out_excel,
    "settings": {"mode": "auto", "singleSheet": False}
})
check("pdf2excel call succeeded", r.get("success"), r.get("error"))
check("pdf2excel output .xlsx exists", os.path.exists(out_excel))
if os.path.exists(out_excel):
    check("pdf2excel output is non-empty", os.path.getsize(out_excel) > 0)

# -----------------------------------------------------------------
# TEST 9: doc_convert.py  pdf2ppt
# -----------------------------------------------------------------
section("TEST 9: pdf2ppt with slideSize=4:3")
out_ppt = os.path.join(tmp_dir, "out.pptx")
r = run("doc_convert.py", {
    "action": "pdf2ppt",
    "file": MINI_PDF,
    "output": out_ppt,
    "settings": {"slideSize": "4:3", "vectorMode": False}
})
check("pdf2ppt call succeeded", r.get("success"), r.get("error"))
check("pdf2ppt output .pptx exists", os.path.exists(out_ppt))
if os.path.exists(out_ppt):
    check("pdf2ppt output is non-empty", os.path.getsize(out_ppt) > 0)

# -----------------------------------------------------------------
# TEST 10: basic_manipulation.py  merge + split + rotate
# -----------------------------------------------------------------
section("TEST 10: basic_manipulation - merge, split, rotate (regression)")
out_merged = os.path.join(tmp_dir, "merged.pdf")
r = run("basic_manipulation.py", {
    "action": "merge",
    "files": [MINI_PDF, MINI_PDF],
    "output": out_merged,
    "add_blank_page": False,
    "compress": False
})
check("merge call succeeded", r.get("success"), r.get("error"))
check("merge output exists", os.path.exists(out_merged))

out_rotated = os.path.join(tmp_dir, "rotated.pdf")
r = run("basic_manipulation.py", {
    "action": "rotate",
    "file": MINI_PDF,
    "output": out_rotated,
    "degrees": 90
})
check("rotate call succeeded", r.get("success"), r.get("error"))
check("rotate output exists", os.path.exists(out_rotated))

out_split_dir = os.path.join(tmp_dir, "split_out")
os.makedirs(out_split_dir, exist_ok=True)
r = run("basic_manipulation.py", {
    "action": "split",
    "file": MINI_PDF,
    "output_dir": out_split_dir,
    "split_mode": "all",
    "range": ""
})
check("split call succeeded", r.get("success"), r.get("error"))

# -----------------------------------------------------------------
# SUMMARY
# -----------------------------------------------------------------
print(f"\n{'='*60}")
print(f"  RESULTS: {PASS} PASSED  |  {FAIL} FAILED")
print(f"{'='*60}\n")

# Cleanup
shutil.rmtree(tmp_dir, ignore_errors=True)

sys.exit(0 if FAIL == 0 else 1)
